from docx import Document
from fastapi.responses import StreamingResponse
import io

def generate_word_response(question_text: str, filename="output.docx"):
    doc = Document()
    doc.add_heading("题目", level=1)
    doc.add_paragraph(question_text)

    # 写入内存而不是文件
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        content=buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=question.docx"}
    )